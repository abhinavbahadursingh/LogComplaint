from email import policy
from email.parser import BytesParser
from io import BytesIO


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Pull plain text out of a complaint document based on file extension."""
    ext = (filename or "").rsplit(".", 1)[-1].lower()
    if ext == "txt":
        return data.decode("utf-8", errors="replace")
    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "eml":
        return _extract_eml(data)
    raise ValueError(f"Unsupported file type: .{ext}")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines).strip()


def _extract_eml(data: bytes) -> str:
    msg = BytesParser(policy=policy.default).parsebytes(data)
    parts = [msg.get("Subject", ""), msg.get("From", ""), msg.get("To", "")]
    body = []
    if msg.is_multipart():
        for part in msg.iter_parts():
            if part.get_content_type() == "text/plain":
                body.append(part.get_content())
    else:
        body.append(msg.get_content())
    return "\n".join([p for p in parts + body if p]).strip()